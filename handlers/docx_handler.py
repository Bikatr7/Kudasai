## Not in use.
##---------------------------------------
## built-in libraries
import zipfile
import io
import typing
import hashlib

## third-party libraries
from docx import Document

import requests

class docxHandler():

    """
    Handles docx to txt conversions.
    """

##--------------------start-of-__init__()------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------

    def __init__(self, inc_docx_path:str, inc_image_kanji_replacement_dict:dict) -> None:

        """

        Initializes the docxHandler class.\n

        Parameters:\n
        inc_docx_path (str) : The path to the docx file.\n
        inc_image_kanji_replacement_dict (dict) : The dictionary containing the image replacements.\n

        Returns:\n
        None.\n

        """

        ## The path to the docx file.
        self.docx_path = inc_docx_path

        ## The dictionary containing the image replacements.
        self.image_kanji_replacement_dict = inc_image_kanji_replacement_dict["image_replacements"]

        ## A dictionary mapping the hash of each known image to its replacement string.
        self.known_images = self.precompute_known_images_from_urls()

##--------------------start-of-download_image()--------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------

    def download_image(self, url:str) -> io.BytesIO:

        """

        Downloads an image from a URL.\n

        Parameters:\n
        url (str) : The URL of the image to download.\n

        Returns:\n
        io.BytesIO : The downloaded image as bytes.\n
        
        """

        response = requests.get(url)

        if(response.status_code == 200):
            return io.BytesIO(response.content)
        
        else:
            raise Exception(f"Image download failed with status code {response.status_code}.")

##--------------------start-of-compute_image_hash()----------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------

    def compute_image_hash(self, image_bytes:io.BytesIO) -> str:

        """

        Computes the SHA-256 hash of an image.\n

        Parameters:\n
        image_bytes (io.BytesIO) : The image as bytes.\n

        Returns:\n
        str : The SHA-256 hash of the image.\n
        
        """

        sha256 = hashlib.sha256()
        sha256.update(image_bytes.getvalue())

        return sha256.hexdigest()

##--------------------start-of-precompute_known_images_from_urls()-------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------

    def precompute_known_images_from_urls(self) -> dict:

        """

        Precomputes the hashes for known images from their URLs.\n

        Parameters:\n
        None.\n

        Returns:\n
        dict : A dictionary mapping the hash of each known image to its replacement string.\n

        """

        known_images = {}

        for url, replacement_string in self.image_kanji_replacement_dict.items():
            image_bytes = self.download_image(url)
            if(image_bytes):
                image_hash = self.compute_image_hash(image_bytes)
                known_images[image_hash] = replacement_string
        return known_images

##--------------------start-of-image_to_string()------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------

    def image_to_string(self, image_bytes:io.BytesIO) -> str:

        """

        Converts an image to its replacement string based on its hash.\n

        Parameters:\n
        image_bytes (io.BytesIO) : The image as bytes.\n

        Returns:\n
        str : The replacement string for the image, or None if the image is not known.\n

        """

        image_hash = self.compute_image_hash(image_bytes)

        return self.known_images.get(image_hash, None)

##--------------------start-of-process_docx()---------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------

    def process_docx(self, output_txt_path:str) -> list:

        """

        Processes the docx file and replaces known images with their replacement strings.\n

        Parameters:\n
        output_txt_path (str) : The path where the extracted text will be saved as a txt file.\n

        Returns:\n
        list : A list of unmatched images, or None if all images are matched.\n

        """

        extracted_text = []
        unmatched_images = []

        document = Document(self.docx_path)

        with zipfile.ZipFile(self.docx_path, 'r') as docx_zip:

            for para in document.paragraphs:
                run_texts = []
                for run in para.runs:
                    if('Drawing' in run._element.xml):
                        for zip_info in docx_zip.infolist():
                            if(zip_info.filename.startswith('word/media/')):
                                with docx_zip.open(zip_info.filename) as image_stream:

                                    image_bytes = io.BytesIO(image_stream.read())
                                    image_string = self.image_to_string(image_bytes)

                                    if(image_string):
                                        run_texts.append(image_string)
                                    else:
                                        unmatched_images.append("Unmatched image found")
                    else:
                        run_texts.append(run.text)

                extracted_text.append(' '.join(run_texts))

            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        extracted_text.append(cell.text)

        with open(output_txt_path, 'w') as txt_file:
            txt_file.write('\n'.join(extracted_text))

        return unmatched_images if unmatched_images else []